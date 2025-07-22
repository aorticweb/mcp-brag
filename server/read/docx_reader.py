from typing import Iterator

from docx import Document

from server.constants import CHUNK_CHARACTER_LIMIT
from server.read.reader import Reader, SourceType, TextChunk
from server.read.text_reader import _split_text_chunk


class DocxReader(Reader):
    file_path: str
    chunk_size_max: int

    def __init__(self, file_path: str, chunk_size_max: int = CHUNK_CHARACTER_LIMIT.value):
        self.file_path = file_path
        self.chunk_size_max = chunk_size_max

    def source_type(self) -> SourceType:
        return SourceType.LOCAL_DOCX_FILE

    def read(self) -> str:
        """
        Read the DOCX file and return all text content as a single string.

        Returns:
            str: The extracted text content from all paragraphs
        """
        doc = Document(self.file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    def read_iter(self) -> Iterator[TextChunk]:
        """
        Read the DOCX file and yield TextChunk objects for each paragraph.

        Each paragraph becomes a TextChunk with start_index and end_index representing
        the character positions in the concatenated text from all paragraphs.

        Yields:
            TextChunk: Chunk containing paragraph text and its character indices
        """
        doc = Document(self.file_path)
        char_index = 0

        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text

            # Add newline to match the behavior of the read() method
            paragraph_with_newline = paragraph_text + "\n"

            # Only process non-empty paragraphs (after stripping)
            if paragraph_text.strip():
                # Create TextChunk for this paragraph
                paragraph_chunk = TextChunk(
                    start_index=char_index,
                    end_index=char_index + len(paragraph_with_newline),
                    text=paragraph_text.strip(),
                )

                # Split the paragraph chunk if it exceeds the size limit
                for chunk in _split_text_chunk(self.chunk_size_max, paragraph_chunk):
                    yield chunk

            # Always advance the character index by the full paragraph length including newline
            char_index += len(paragraph_with_newline)
